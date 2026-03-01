"""Generate resume as Word document and PDF.

Outputs:
- Alberto_Nieto_Resume.docx  (project root)
- src/assets/Alberto_Nieto_Resume.pdf  (served by the site)

Formatting follows ATS best practices:
- Single column, no tables/text boxes
- Calibri throughout, navy accent (#1B3A5C)
- 0.75in margins, 1.15 line spacing
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import subprocess

# -- Colors --
NAVY = RGBColor(0x1B, 0x3A, 0x5C)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# -- Base style --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = DARK_GRAY
style.paragraph_format.space_after = Pt(1)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.1


def add_bottom_border(paragraph, color="1B3A5C"):
    """Add a thin bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(text):
    """Add a section heading with navy bottom border."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.space_before = Pt(10)
    p.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = NAVY
    add_bottom_border(p)
    return p


def add_role(title, company_location, period, summary="", bullets=None):
    """Add a role entry with title, company, period, summary, and bullets."""
    # Title + right-aligned period
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(10.5)
    run_title.font.name = "Calibri"
    run_title.font.color.rgb = DARK_GRAY
    p.add_run("\t")
    run_period = p.add_run(period)
    run_period.font.size = Pt(9.5)
    run_period.font.name = "Calibri"
    run_period.font.color.rgb = MID_GRAY
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.6), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Company line
    p2 = doc.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(1)
    p2.paragraph_format.keep_with_next = True
    run_co = p2.add_run(company_location)
    run_co.font.size = Pt(9.5)
    run_co.font.name = "Calibri"
    run_co.font.color.rgb = MID_GRAY

    if summary:
        p3 = doc.add_paragraph()
        p3.space_before = Pt(0)
        p3.space_after = Pt(1)
        p3.paragraph_format.keep_with_next = True
        run_s = p3.add_run(summary)
        run_s.font.size = Pt(10)
        run_s.font.name = "Calibri"
        run_s.font.color.rgb = DARK_GRAY

    if bullets:
        for bullet in bullets:
            bp = doc.add_paragraph(style="List Bullet")
            bp.clear()
            bp.space_before = Pt(0)
            bp.space_after = Pt(1)
            bp.paragraph_format.left_indent = Inches(0.25)
            run_b = bp.add_run(bullet)
            run_b.font.size = Pt(10)
            run_b.font.name = "Calibri"
            run_b.font.color.rgb = DARK_GRAY


# ===== HEADER =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(2)
run = p.add_run("ALBERTO NIETO")
run.bold = True
run.font.size = Pt(24)
run.font.name = "Calibri"
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(0)
run = p.add_run(
    "Alberto.Codes.Dev@gmail.com  |  alberto.codes  |  github.com/Alberto-Codes"
)
run.font.size = Pt(10)
run.font.name = "Calibri"
run.font.color.rgb = NAVY

# ===== PROFESSIONAL SUMMARY =====
add_section_heading("Professional Summary")
p = doc.add_paragraph()
p.space_after = Pt(4)
run = p.add_run(
    "Strategic and hands-on technology leader with 25+ years in financial services, "
    "specializing in AI-driven automation, cloud-native solutions, and enterprise-scale "
    "data engineering. Career spans from banking operations to Principal Engineer \u2014 "
    "building deep domain expertise at every level. Expert in Agentic AI frameworks, "
    "GCP infrastructure, and scalable automation pipelines. Open source contributor "
    "with published Python libraries on PyPI. Co-inventor on patent application. "
    "Bilingual in English and Spanish."
)
run.font.size = Pt(10.5)
run.font.name = "Calibri"
run.font.color.rgb = DARK_GRAY

# ===== AREAS OF EXPERTISE =====
add_section_heading("Areas of Expertise")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(4)
run = p.add_run(
    "Generative AI & Agents  \u2022  Python  \u2022  Cloud Engineering (GCP, PCF)  \u2022  "
    "Data Engineering & Pipelines\n"
    "OCR & Document Intelligence  \u2022  Video Inference  \u2022  "
    "CI/CD & DevOps  \u2022  Enterprise Architecture"
)
run.font.size = Pt(10.5)
run.font.name = "Calibri"
run.font.color.rgb = DARK_GRAY

# ===== OPEN SOURCE =====
add_section_heading("Open Source")
for name, desc in [
    (
        "adk-secure-sessions",
        "Encrypted session storage for Google ADK agents. "
        "Drop-in replacement using Fernet encryption with PBKDF2 key derivation.",
    ),
    (
        "docvet",
        "CLI tool for Python docstring quality vetting. 19 rules across "
        "completeness, accuracy, rendering, and visibility.",
    ),
    (
        "gepa-adk",
        "Evolutionary prompt optimization for AI agents using Google ADK. "
        "Automatically evolves agent instructions through iterative improvement.",
    ),
]:
    bp = doc.add_paragraph()
    bp.space_before = Pt(1)
    bp.space_after = Pt(2)
    bp.paragraph_format.left_indent = Inches(0.25)
    run_name = bp.add_run(f"{name}: ")
    run_name.bold = True
    run_name.font.size = Pt(10.5)
    run_name.font.name = "Calibri"
    run_name.font.color.rgb = NAVY
    run_desc = bp.add_run(desc)
    run_desc.font.size = Pt(10.5)
    run_desc.font.name = "Calibri"
    run_desc.font.color.rgb = DARK_GRAY

# ===== RECOGNITION =====
add_section_heading("Recognition")
for item in [
    "January 2026 \u2014 Named as co-inventor on patent application",
    "Top Performer 2018 (New Orleans) \u2014 Recognition for vulnerability remediation efforts",
    "Top Performer 2014 (Nashville) \u2014 Recognition for asset reporting initiatives",
]:
    bp = doc.add_paragraph(style="List Bullet")
    bp.clear()
    bp.space_before = Pt(0)
    bp.space_after = Pt(2)
    bp.paragraph_format.left_indent = Inches(0.25)
    run_b = bp.add_run(item)
    run_b.font.size = Pt(10.5)
    run_b.font.name = "Calibri"
    run_b.font.color.rgb = DARK_GRAY

# ===== PROFESSIONAL EXPERIENCE =====
add_section_heading("Professional Experience")

add_role(
    "Principal Engineer (Executive Director)",
    "Wells Fargo, Phoenix, Arizona",
    "Jun 2025 \u2013 Present",
    "Leader in implementing cutting-edge AI solutions for business needs, "
    "presented to CTO weekly. Creating documentation and frameworks to "
    "accelerate developer onboarding and production rollouts.",
    [
        "Created AI agent framework to automate line-of-business operations, "
        "adopted across multiple teams",
        "Built video processing pipeline handling thousands of videos daily into "
        "standard operating procedures for knowledge tree persistence",
        "Implemented evaluation and evolution frameworks to score and evolve "
        "agent definitions for data pipeline tasks (open-sourced as gepa-adk)",
        "Established patterns ensuring non-functional requirements are satisfied "
        "with AI deployments",
        "Audited enterprise API clients for Apigee AI provisioning to enable "
        "multi-modal capabilities and proper scaling",
    ],
)

add_role(
    "Senior Lead Analytics Consultant (Executive Director)",
    "Wells Fargo, Phoenix, Arizona",
    "Jun 2022 \u2013 Jun 2025",
    "Individual contributor and team's resident expert in GenAI and Agentic AI "
    "frameworks. Focused on proof of concepts and production-ready AI implementations.",
    [
        "Architected scalable OCR pipeline processing 500K+ documents using Gemini, "
        "with results stored in Teradata and SQLite for distributed processing",
        "Developed intelligent agents for database querying, KPI summarization, "
        "SQL validation, and structured QA using Python and Google ADK",
        "Built reusable data pipelines for document ingestion, OCR, and validation "
        "using Selenium, HTTPX, and Docling with Gradio/Streamlit frontends",
        "Designed modular, scalable architecture patterns reused across multiple "
        "teams and projects",
        "Implemented git-based dev workflow with peer reviews, pull requests, and "
        "release standards to mature end-user compute tool development",
    ],
)

add_role(
    "Lead Analytics Consultant (AVP)",
    "Wells Fargo, Phoenix, Arizona",
    "Jan 2018 \u2013 Jun 2022",
    "T-shaped role providing leadership, oversight, and direct development "
    "in ETL, visualization, and advanced analytics workstreams.",
    [
        "Onboarded applications to Pivotal Cloud Foundry with Splunk logging "
        "and Blackduck security scans",
        "Built web scrape automation for business intelligence visualization "
        "previews and validation",
        "Created dynamic databases for manual and automated data ingestion "
        "from flat-file data providers",
        "Engaged technology peer groups to execute high-priority regulatory "
        "data analytics",
    ],
)

add_role(
    "Applications Systems Engineer (AVP)",
    "Wells Fargo, Phoenix, Arizona",
    "Nov 2016 \u2013 Jan 2018",
    "Team lead for vulnerability business intelligence development team.",
    [
        "Provided guidance to development team and business partners on BI "
        "solutions for emergency remediation and reporting",
        "Consulted leadership on data architecture roadmap for vulnerability "
        "remediation",
        "Created ad-hoc SQL reports across systems of record to illustrate "
        "compliance status",
    ],
)

add_role(
    "Business Systems Consultant (AVP)",
    "Wells Fargo, Phoenix, Arizona",
    "Oct 2009 \u2013 Nov 2016",
    "Created business intelligence reporting and tools for regulatory compliance "
    "and leadership decision-making.",
    [
        "Built and supported ongoing asset validation processes",
        "Created compliance metrics for server footprint tracking",
        "Developed business process reports for server implementation, changes, "
        "and decommission time-to-market",
    ],
)

add_role(
    "Technology Manager",
    "Wells Fargo, Phoenix, Arizona",
    "2006 \u2013 2009",
    "Led team of technical professionals supporting end-user technology across "
    "four administrative sites.",
    [
        "Led 8-9 technicians across four administrative sites",
        "Led desktop OS migrations and provided onsite hardware swap support",
        "Expedited Business Continuity Plan documentation to 100% compliance "
        "for disaster recovery across four business lines",
        "Led monthly compliance efforts for security patching, encryption, "
        "and asset tracking",
    ],
)

add_role(
    "Team Lead / Leadership Development Program / Operations",
    "Wells Fargo / Bank of America, Phoenix, Arizona",
    "1999 \u2013 2006",
    "Progressive career from teller to operations analyst to leadership "
    "development program, building deep financial services domain expertise.",
    [
        "Selected for Wells Fargo Leadership Development Program (2004-2005), "
        "rotating through operations and technology departments",
        "Presented SharePoint hosting solution to technology executive leadership",
        "Led six technicians at two admin sites with desktop end-user support",
        "Progressed through teller, personal banker, lead teller, operations "
        "processor, and operations analyst roles (1999-2004)",
    ],
)

# ===== EDUCATION =====
add_section_heading("Education")
p = doc.add_paragraph()
p.space_before = Pt(2)
run_deg = p.add_run("Bachelor of Science, Accounting")
run_deg.bold = True
run_deg.font.size = Pt(11)
run_deg.font.name = "Calibri"
run_deg.font.color.rgb = DARK_GRAY
p2 = doc.add_paragraph()
p2.space_before = Pt(0)
run_school = p2.add_run("DeVry University, Phoenix, Arizona")
run_school.font.size = Pt(10.5)
run_school.font.name = "Calibri"
run_school.font.color.rgb = MID_GRAY

# ===== SAVE =====
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

docx_path = os.path.join(project_root, "Alberto_Nieto_Resume.docx")
doc.save(docx_path)
print(f"DOCX saved to: {docx_path}")

# Convert to PDF via LibreOffice
pdf_dest = os.path.join(project_root, "src", "assets", "Alberto_Nieto_Resume.pdf")
try:
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            os.path.dirname(pdf_dest),
            docx_path,
        ],
        check=True,
        capture_output=True,
    )
    print(f"PDF saved to: {pdf_dest}")
except FileNotFoundError:
    print("LibreOffice not found -- DOCX saved but PDF not generated.")
    print(f"To convert manually: libreoffice --headless --convert-to pdf {docx_path}")
